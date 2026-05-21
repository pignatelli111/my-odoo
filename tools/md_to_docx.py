#!/usr/bin/env python3
"""Minimal Markdown → DOCX converter for SBU docs (headings, tables, lists, code)."""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt


def strip_md_inline(text: str) -> str:
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    return text.strip()


def add_runs_with_bold(paragraph, text: str):
    parts = re.split(r'(\*\*.+?\*\*)', text)
    for part in parts:
        if not part:
            continue
        m = re.match(r'\*\*(.+?)\*\*', part)
        if m:
            run = paragraph.add_run(m.group(1))
            run.bold = True
        else:
            plain = strip_md_inline(part)
            if plain:
                paragraph.add_run(plain)


def parse_table_row(line: str) -> list[str]:
    line = line.strip().strip('|')
    return [strip_md_inline(c.strip()) for c in line.split('|')]


def is_table_sep(line: str) -> bool:
    return bool(re.match(r'^\|[\s\-:|]+\|$', line.strip()))


def convert(md_path: Path, docx_path: Path):
    lines = md_path.read_text(encoding='utf-8').splitlines()
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    i = 0
    in_code = False
    code_lines: list[str] = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith('```'):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if stripped == '---':
            i += 1
            continue

        if stripped.startswith('# '):
            doc.add_heading(strip_md_inline(stripped[2:]), level=0)
            i += 1
            continue
        if stripped.startswith('## '):
            doc.add_heading(strip_md_inline(stripped[3:]), level=1)
            i += 1
            continue
        if stripped.startswith('### '):
            doc.add_heading(strip_md_inline(stripped[4:]), level=2)
            i += 1
            continue

        if stripped.startswith('|') and i + 1 < len(lines) and is_table_sep(lines[i + 1]):
            headers = parse_table_row(stripped)
            i += 2
            rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                rows.append(parse_table_row(lines[i]))
                i += 1
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = 'Table Grid'
            for c, h in enumerate(headers):
                table.rows[0].cells[c].text = h
            for r, row in enumerate(rows):
                for c in range(len(headers)):
                    table.rows[r + 1].cells[c].text = row[c] if c < len(row) else ''
            doc.add_paragraph()
            continue

        if stripped.startswith('- [ ] '):
            p = doc.add_paragraph(style='List Bullet')
            add_runs_with_bold(p, '☐ ' + stripped[6:])
            i += 1
            continue
        if stripped.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            add_runs_with_bold(p, stripped[2:])
            i += 1
            continue

        m = re.match(r'^(\d+)\.\s+(.+)$', stripped)
        if m:
            p = doc.add_paragraph(style='List Number')
            add_runs_with_bold(p, m.group(2))
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        p = doc.add_paragraph()
        add_runs_with_bold(p, stripped)
        i += 1

    doc.save(docx_path)
    print(f'Wrote {docx_path}')


if __name__ == '__main__':
    src = Path(sys.argv[1]) if len(sys.argv) > 1 else Path('docs/FEEDBACK_COSIMO_ROADMAP_IT.md')
    dst = Path(sys.argv[2]) if len(sys.argv) > 2 else src.with_suffix('.docx')
    convert(src.resolve(), dst.resolve())
